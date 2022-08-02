

from itertools import groupby
from heapq import *
import os

from bitstring import BitArray, BitStream
from bitarray import bitarray

from app.utils import *
from docx import *


class Node:
	left = None
	right = None
	item = None
	value = 0

	def __init__(self, item, value):
		self.item = item
		self.value = value

	def set_children(self, left_node, right_node):
		self.left = left_node
		self.right = right_node

	def __lt__(self, other):
		return self.value < other.value



class HuffmanCoder:

	def __init__(self, filepath, destination_path):
		self.filepath = filepath
		self.destination_path = destination_path

		self.filename = self.filepath.split("/")[-1].split(".")[0]
		self.file_ext = os.path.splitext(filepath)[1]
		self.bit_length = 0

		self.content = ""

		self.char_to_code = {}
		self.code_to_char = {}

		if self.file_ext == ".txt":
			with open(self.filepath, 'r+') as file:
				self.content = file.read()

		elif self.file_ext == ".docx":
			document = Document(filepath)

			for para in document.paragraphs:
				self.content += para.text

		else:
			raise ValueError("Wrong File Type for Huffman Coding.")
			#return

		self.freq_tree = self.build_freq_tree(self.content)



	def build_freq_tree(self, content):
		freq = {}
		#
		for char in content:
			if not char in freq:
				freq[char] = 0

			freq[char] += 1

		return freq


	def encode(self):
		codes = {}

		def codeIt(s, node):
			#print(node.item)
			#print(s)
			if node.item:
				if not s:
					codes[node.item] = "0"
				else:
					codes[node.item] = s
			else:
				codeIt(s+"0", node.left)
				codeIt(s+"1", node.right)
		#print(sorted(self.content))


		queue = [Node(a, len(list(b))) for a, b in groupby(sorted(self.content))]
		#print(a)
		#print(len(list(b)))
		#print(queue)
		heapify(queue)
		#print(queue)


		while len(queue) > 1:
			left = heappop(queue)
			right = heappop(queue)

			node = Node(None, right.value + left.value)
			node.set_children(left, right)
			heappush(queue, node)

		codeIt("", queue[0])
		#print(codes)
		return codes


	def compress(self):
		compressed_filepath = self.destination_path + "/" + self.filename + "-" + self.file_ext[1:] + ".bin"

		# Coding for .txt and .docx files.
		if self.file_ext == ".txt" or ".docx":

			self.char_to_code = self.encode()
			self.code_to_char = {value:key for key, value in self.char_to_code.items()}
			encoded_content = "".join([self.char_to_code[a] for a in self.content])
			#print(encoded_content)
			self.bit_length = len(encoded_content)
			bit_array = BitArray(bin = encoded_content)
			#print(bit_array)

			with open(compressed_filepath, 'wb') as compressed_file:
				compressed_file.write(bit_array.tobytes())
				compressed_file.close()

			return compressed_filepath

		else:
			raise ValueError("HuffmanCoder Compression Error: Wrong File Type")
			return



	def decompress(self, compressed_filepath, decompression_path):
		bits = BitArray(filename = compressed_filepath)

		encoded_content = bits.bin[0:self.bit_length]
		decoder = {k:bitarray(v) for k, v in self.char_to_code.items()}
		#print(decoder)
		decoded_chars = bitarray(encoded_content).decode(decoder)
		decoded_content = ''.join(x for x in decoded_chars)

		if self.file_ext == ".txt":
			decompressed_filepath = decompression_path + "/" + self.filename + "_(decompressed)" + ".txt"

			with open(decompressed_filepath, 'w+') as decompressed_file:
				decompressed_file.write(decoded_content)
				decompressed_file.close()

			return decompressed_filepath

		if self.file_ext == ".docx":
			decompressed_filepath = decompression_path + "/" + self.filename + "_(decompressed)" + ".docx"

			document = Document()
			document.add_paragraph(decoded_content)
			document.save(decompressed_filepath)

			return decompressed_filepath



	#def print_attrs(self):
		#print("Filepath: " + self.filepath)
		#print("File Name: " + self.filename)
		#print("File Extension: " + self.file_ext)
		#print(self.freq_tree)

